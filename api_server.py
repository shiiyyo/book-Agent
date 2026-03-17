# -*- coding: utf-8 -*-
"""
极简后端：只提供 API 给前端调用；生成任务在后台线程执行，支持实时取消与微调（编辑保存、修改说明）。
启动：在项目根目录执行  py api_server.py  ，浏览器打开  http://127.0.0.1:5000
"""
from __future__ import annotations

import io
import json
import re
import sys
import threading
from datetime import datetime, timedelta, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

from dotenv import load_dotenv
load_dotenv(ROOT / ".env")

from flask import Flask, Response, jsonify, request, send_file, send_from_directory, stream_with_context
from sqlalchemy import delete, select, text

from app.config import get_default_model
from app.pipeline.task_runner import run_generation_task
from core.engine import _build_outline_content, _parse_outline_to_intro_and_fragments
from app.semantic.citation_manager import add_reference_from_file, add_reference_from_url, list_references
from database import (
    Argument,
    Book,
    Chapter,
    Citation,
    Conflict,
    GenerationTask,
    Outline,
    Reference,
    TaskStatus,
    TaskType,
    Term,
    get_session,
    init_db,
)
from database.connection import ensure_content_type_column


def _run_task_in_background(book_id: int, task_id: int) -> None:
    """在后台线程中执行单条生成任务，便于前端轮询与取消。"""
    session = get_session()
    try:
        task = session.get(GenerationTask, task_id)
        if not task or task.book_id != book_id or task.status != TaskStatus.PENDING:
            return
        task.status = TaskStatus.RUNNING
        task.progress_message = "运行中…"
        task.started_at = datetime.now(timezone.utc).replace(tzinfo=None)
        task.completed_at = None
        session.commit()
        session.refresh(task)
        glossary = _load_glossary(session, book_id) if task.task_type in (TaskType.CHAPTER, TaskType.REWRITE) else None
        run_generation_task(session, task, glossary_terms=glossary)
        session.refresh(task)
        if task.status != TaskStatus.CANCELLED:
            session.commit()
        else:
            session.rollback()
    except Exception as e:
        session.rollback()
        try:
            task = session.get(GenerationTask, task_id)
            if task:
                task.status = TaskStatus.FAILED
                task.error_message = (str(e) or str(sys.exc_info()[1]))[:2000]
                session.commit()
        except Exception:
            pass
    finally:
        session.close()

app = Flask(__name__, static_folder=str(ROOT / "static"), static_url_path="")
app.config["JSON_AS_ASCII"] = False

# 参考文献上传保存目录
REF_DIR = ROOT / "data" / "references"
REF_DIR.mkdir(parents=True, exist_ok=True)


def _cleanup_stale_running_tasks() -> int:
    """
    服务器重启后后台线程不会自动续跑。
    为避免 UI 永远显示“运行中”，将遗留的 RUNNING 任务标记为 FAILED。
    """
    session = get_session()
    try:
        running = list(session.execute(
            select(GenerationTask).where(GenerationTask.status == TaskStatus.RUNNING)
        ).scalars().all())
        if not running:
            return 0
        now = datetime.now(timezone.utc).replace(tzinfo=None)
        for t in running:
            t.status = TaskStatus.FAILED
            t.error_message = (t.error_message or "") + ("\n" if (t.error_message or "").strip() else "") + "服务已重启：该任务在重启时被中断并标记为失败。"
            t.progress_message = "已中断（服务重启）"
            t.completed_at = now
        session.commit()
        return len(running)
    except Exception:
        session.rollback()
        return 0
    finally:
        session.close()


@app.route("/")
def index():
    return send_from_directory(app.static_folder, "index.html")


# ---------- API ----------

@app.route("/api/config", methods=["GET"])
def get_config():
    """返回当前配置（如调用模型来自 .env），供前端展示。"""
    return jsonify({"default_model": get_default_model()})


@app.route("/api/books", methods=["GET"])
def list_books():
    session = get_session()
    try:
        books = session.execute(select(Book).order_by(Book.updated_at.desc())).scalars().all()
        return jsonify([{"id": b.id, "title": b.title, "core_concept": b.core_concept, "default_model": b.default_model or "", "content_type": getattr(b, "content_type", None) or "academic"} for b in books])
    finally:
        session.close()


@app.route("/api/books", methods=["POST"])
def create_book():
    data = request.get_json() or {}
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "书名为空"}), 400
    content_type = (data.get("content_type") or "academic").strip().lower()
    if content_type not in ("academic", "novel"):
        content_type = "academic"
    session = get_session()
    try:
        default_model = (data.get("default_model") or get_default_model() or "").strip() or None
        book = Book(title=title, core_concept=(data.get("core_concept") or "").strip() or None, default_model=default_model, content_type=content_type, book_type=content_type)
        session.add(book)
        session.commit()
        session.refresh(book)
        return jsonify({"id": book.id, "title": book.title, "content_type": getattr(book, "content_type", None) or "academic"})
    finally:
        session.close()


@app.route("/api/books/<int:bid>/export", methods=["GET"])
def export_book(bid):
    """导出书籍为 Markdown 或 Word。?format=md 或 format=docx（须在 get_book 之前注册以免被覆盖）"""
    fmt = (request.args.get("format") or "md").strip().lower()
    if fmt not in ("md", "docx"):
        return jsonify({"error": "仅支持 format=md 或 format=docx"}), 400
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        outline = book.outline
        chapters = sorted(book.chapters, key=lambda c: c.order_index)
        base_name = _sanitize_filename(book.title or "book")
        if fmt == "md":
            content = _book_to_markdown(book, outline, chapters)
            buf = io.BytesIO(content.encode("utf-8"))
            return send_file(
                buf,
                mimetype="text/markdown; charset=utf-8",
                as_attachment=True,
                download_name=base_name + ".md",
            )
        else:
            try:
                raw = _book_to_docx(book, outline, chapters)
            except ImportError:
                return jsonify({"error": "请安装 python-docx: pip install python-docx"}), 500
            return send_file(
                io.BytesIO(raw),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=base_name + ".docx",
            )
    finally:
        session.close()


@app.route("/api/books/<int:bid>", methods=["GET"])
def get_book(bid):
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        outline = None
        chapters = sorted(book.chapters, key=lambda c: c.order_index)
        if book.outline:
            content = book.outline.content or ""
            intro = getattr(book.outline, "intro", None) or ""
            if intro is not None and any(getattr(c, "outline_fragment", None) and (c.outline_fragment or "").strip() for c in chapters):
                content = _build_outline_content(intro, chapters) or content
            outline = {"content": content, "raw_json": book.outline.raw_json}
        terms = [{"term": t.term, "definition": t.definition or ""} for t in book.terms]
        return jsonify({
            "id": book.id,
            "title": book.title,
            "core_concept": book.core_concept,
            "default_model": book.default_model or "",
            "content_type": getattr(book, "content_type", None) or "academic",
            "preface": getattr(book, "preface", None) or "",
            "target_publisher": getattr(book, "target_publisher", None) or "",
            "writing_style": getattr(book, "writing_style", None) or "",
            "style_reference_text": getattr(book, "style_reference_text", None) or "",
            "academic_tone": getattr(book, "academic_tone", None) or "strict",
            "outline": outline,
            "chapters": [{
                "id": c.id,
                "order_index": c.order_index,
                "title": c.title,
                "outline_content": c.outline_content,
                "draft_content": getattr(c, "draft_content", None),
                "approved_sections": getattr(c, "approved_sections", None),
                "content": c.content,
                "outline_fragment": getattr(c, "outline_fragment", None) or "",
            } for c in chapters],
            "terms": terms,
        })
    finally:
        session.close()


@app.route("/api/books/<int:bid>/check", methods=["GET"])
def check_book_content_type(bid):
    """诊断接口：确保 content_type 列存在，并返回该书在数据库中的 content_type 实际值。"""
    ensure_content_type_column()
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        orm_value = getattr(book, "content_type", None)
        raw_value = None
        try:
            from database.connection import engine
            with engine.connect() as conn:
                row = conn.execute(text("SELECT content_type FROM books WHERE id = :id"), {"id": bid}).fetchone()
                raw_value = row[0] if row else None
        except Exception as e:
            raw_value = f"(查询失败: {e})"
        return jsonify({
            "book_id": bid,
            "title": book.title,
            "content_type_orm": orm_value,
            "content_type_raw": raw_value,
            "hint": "若 content_type_raw 为 null 或 academic 但您选的是网络小说，请在前端把内容类型改为「网络小说」并保存后再生成。",
        })
    finally:
        session.close()


@app.route("/api/books/<int:bid>", methods=["PUT"])
def update_book(bid):
    data = request.get_json() or {}
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        if "default_model" in data:
            book.default_model = (data["default_model"] or "").strip() or book.default_model
        if "content_type" in data and hasattr(book, "content_type"):
            ct = (data["content_type"] or "").strip().lower()
            if ct in ("academic", "novel"):
                book.content_type = ct
                if hasattr(book, "book_type"):
                    book.book_type = ct
        if "preface" in data and hasattr(book, "preface"):
            if data["preface"] is not None:
                book.preface = _normalize_paragraph_spacing(data["preface"])
        if "target_publisher" in data and hasattr(book, "target_publisher"):
            book.target_publisher = (data["target_publisher"] or "").strip() or None
        if "writing_style" in data and hasattr(book, "writing_style"):
            book.writing_style = (data["writing_style"] or "").strip() or None
        if "style_reference_text" in data and hasattr(book, "style_reference_text"):
            book.style_reference_text = (data["style_reference_text"] or "").strip() or None
        if "academic_tone" in data and hasattr(book, "academic_tone"):
            t = (data["academic_tone"] or "").strip().lower()
            book.academic_tone = t if t in ("strict", "bestseller") else (book.academic_tone or "strict")
        session.commit()
        return jsonify({"ok": True})
    finally:
        session.close()


@app.route("/api/books/<int:bid>/outline", methods=["PUT"])
def update_outline(bid):
    """保存用户对大纲内容的手动修改（微调）；同步拆分为 intro + 各章 outline_fragment 便于局部修改。"""
    data = request.get_json() or {}
    content = (data.get("content") or "").strip()
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        content = _normalize_paragraph_spacing(content or "")
        if not book.outline:
            from database.models import Outline
            book.outline = Outline(book_id=bid, content=content or "", raw_json=None)
            session.add(book.outline)
        else:
            book.outline.content = content or book.outline.content or ""
        intro, fragments = _parse_outline_to_intro_and_fragments(content or "")
        if getattr(book.outline, "intro", None) is not None:
            book.outline.intro = intro
        chapters = sorted(book.chapters, key=lambda c: c.order_index)
        frag_by_idx = {idx: frag for idx, frag in fragments}
        for ch in chapters:
            if ch.order_index in frag_by_idx and hasattr(ch, "outline_fragment"):
                ch.outline_fragment = frag_by_idx[ch.order_index]
        session.commit()
        return jsonify({"ok": True})
    finally:
        session.close()


@app.route("/api/books/<int:bid>/chapters/<int:cid>", methods=["PUT"])
def update_chapter(bid, cid):
    """保存用户对章节正文的手动修改（微调）。"""
    data = request.get_json() or {}
    if "content" in data:
        session = get_session()
        try:
            book = session.get(Book, bid)
            ch = session.get(Chapter, cid)
            if not book or not ch or ch.book_id != bid:
                return jsonify({"error": "书籍或章节不存在"}), 404
            raw_content = data.get("content")
            if raw_content is not None:
                ch.content = _normalize_paragraph_spacing(raw_content)
            if ch.content:
                ch.word_count = len(ch.content)
            session.commit()
            return jsonify({"ok": True})
        finally:
            session.close()
    return jsonify({"error": "请提供 content 字段"}), 400


@app.route("/api/books/<int:bid>", methods=["DELETE"])
def delete_book(bid):
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        # 按依赖顺序删除关联数据，避免 SQLite 外键或 CASCADE 未生效时报错
        session.execute(delete(Conflict).where(Conflict.book_id == bid))
        session.execute(delete(GenerationTask).where(GenerationTask.book_id == bid))
        chapter_ids = [r[0] for r in session.execute(select(Chapter.id).where(Chapter.book_id == bid)).all()]
        if chapter_ids:
            session.execute(delete(Citation).where(Citation.chapter_id.in_(chapter_ids)))
        session.execute(delete(Chapter).where(Chapter.book_id == bid))
        session.execute(delete(Reference).where(Reference.book_id == bid))
        session.execute(delete(Argument).where(Argument.book_id == bid))
        session.execute(delete(Term).where(Term.book_id == bid))
        session.execute(delete(Outline).where(Outline.book_id == bid))
        session.delete(book)
        session.commit()
        return jsonify({"ok": True})
    except Exception as e:
        session.rollback()
        return jsonify({"error": "删除失败: " + str(e)}), 500
    finally:
        session.close()


def _sanitize_filename(name: str, max_len: int = 80) -> str:
    """将书名转为安全文件名：去掉非法字符并截断长度。"""
    if not name or not name.strip():
        return "book"
    s = re.sub(r'[\\/:*?"<>|]', "_", name.strip())
    return s[:max_len] if len(s) > max_len else s


def _book_to_markdown(book, outline, chapters) -> str:
    """将书籍内容拼接为纯文本导出（已去除 *、# 与多余空行，无 Markdown 符号）。"""
    def plain(s):
        return _export_plain(s) if s else ""
    lines = [plain(book.title or "未命名"), ""]
    if book.core_concept:
        lines.append("核心构思\n\n" + plain(book.core_concept or "") + "\n\n---\n")
    preface = getattr(book, "preface", None) or ""
    if preface and preface.strip():
        lines.append("前言\n\n" + plain(preface) + "\n\n---\n")
    if outline and outline.content:
        lines.append("全书大纲\n\n" + plain(outline.content or "") + "\n\n---\n")
    for ch in chapters:
        lines.append("第 {} 章 {}".format(ch.order_index, plain(ch.title or "")))
        if ch.outline_content and (ch.outline_content or "").strip():
            lines.append("\n本章要点： " + plain(ch.outline_content or "") + "\n")
        ap = getattr(ch, "approved_sections", None)
        if isinstance(ap, str):
            try:
                ap = json.loads(ap) if ap.strip() else None
            except Exception:
                ap = None
        items = ap.get("items") if isinstance(ap, dict) else None
        body_src = ch.content or ""
        draft_src = getattr(ch, "draft_content", None) or ""
        if items and isinstance(items, dict) and draft_src.strip():
            body_src = draft_src
            for k, v in items.items():
                if k and v and str(v).strip():
                    body_src = _replace_markdown_section_by_title(body_src, str(k), str(v))
        elif not (body_src or "").strip() and draft_src.strip():
            body_src = draft_src
        if body_src and body_src.strip():
            lines.append("\n" + plain(body_src))
        else:
            lines.append("\n（本章正文尚未生成）")
        lines.append("\n")
    return "\n".join(lines)


def _filter_markdown_by_titles(text: str, titles: list[str]) -> str:
    """从章节 Markdown 中仅保留指定的小节（##/### 标题匹配）。"""
    if not text or not titles:
        return text or ""
    keep = set([str(t).strip() for t in titles if t and str(t).strip()])
    if not keep:
        return text or ""
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        m = re.match(r"^(#{2,3})\s+(.+)$", (line or "").strip())
        if not m:
            i += 1
            continue
        title = (m.group(2) or "").strip()
        level = len(m.group(1))
        j = i + 1
        while j < len(lines):
            m2 = re.match(r"^(#{2,3})\s+(.+)$", (lines[j] or "").strip())
            if m2 and len(m2.group(1)) <= level:
                break
            j += 1
        if title in keep:
            out.extend(lines[i:j])
            out.append("")
        i = j
    return "\n".join(out).strip() if out else ""


def _replace_markdown_section_by_title(base_text: str, title: str, new_section_with_heading: str) -> str:
    """在 base_text 中找到指定小节并替换；找不到则追加到末尾。"""
    if not base_text:
        return (new_section_with_heading or "").strip()
    t = (title or "").strip()
    if not t:
        return base_text
    try:
        from core.engine import _split_markdown_section_by_title
        before, _, after = _split_markdown_section_by_title(base_text, t)
        stitched = "\n\n".join([p for p in [before, new_section_with_heading, after] if p and p.strip()]).strip()
        return stitched
    except Exception:
        return ("\n\n".join([base_text.rstrip(), (new_section_with_heading or "").strip()]).strip() if (new_section_with_heading or "").strip() else base_text)


def _normalize_paragraph_spacing(text: str | None) -> str:
    """标准化段落格式：去除 * 号、统一换行、多处空行合并为一段一空行。"""
    if not text or not isinstance(text, str):
        return (text or "").strip()
    import re
    s = text.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("**", "").replace("*", "")
    s = re.sub(r"\n[\s]*\n[\s]*\n+", "\n\n", s)
    s = re.sub(r"\n{2,}", "\n\n", s)
    return s.strip()


def _strip_heading_hashes(text: str | None) -> str:
    """去除每行行首的 Markdown 标题符 #（一个或多个 # 及紧随其后的空白）。"""
    if not text or not isinstance(text, str):
        return (text or "").strip()
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    out = [re.sub(r"^#+\s*", "", line) for line in lines]
    return "\n".join(out).strip()


def _export_plain(text: str | None) -> str:
    """导出用：去除 *、# 与多余空行，得到纯文本。"""
    if not text or not isinstance(text, str):
        return (text or "").strip()
    s = _normalize_paragraph_spacing(text)
    s = _strip_heading_hashes(s)
    return s.strip()


def _strip_markdown_asterisks(text: str) -> str:
    """去除 Markdown 粗体/斜体标记（*、**），导出 Word 时再次确保无星号。"""
    if not text:
        return text
    s = text.replace("**", "")
    s = s.replace("*", "")
    return s


# Word 导出排版规范（与 SPEC.md §7、前端展示一致）：
# 正文：宋体小四、1.5 倍行距、首行缩进 2 字符、两端对齐。
# 一级「第一章」：黑体三号居中。
# 二级「一、」：黑体四号左对齐。
# 三级「（一）」：楷体小四左对齐。
# 强调仅加粗；导出即可直接用于交稿排版。


def _book_to_docx(book, outline, chapters) -> bytes:
    """将书籍内容生成为 Word 文档，遵循统一交稿排版规范（见本函数上方注释与 SPEC.md §7）。"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    def plain(s):
        return _export_plain(s) if s else ""

    doc = Document()
    # 正文：宋体小四、1.5 倍行距、首行缩进 2 字符、两端对齐
    FONT_SONG = "宋体"
    FONT_HEI = "黑体"
    FONT_KAI = "楷体"
    SIZE_XIAOSI = Pt(12)   # 小四
    SIZE_SIHAO = Pt(14)    # 四号
    SIZE_SANHAO = Pt(16)   # 三号
    INDENT_2CH = Cm(0.74)  # 首行缩进约 2 字符

    def set_body_style(p):
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = INDENT_2CH
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for r in p.runs:
            r.font.name = FONT_SONG
            r.font.size = SIZE_XIAOSI
            r.font.bold = False
            r.font.italic = False

    def add_body(doc, text, bold_prefix=None):
        if not text and not bold_prefix:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = INDENT_2CH
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            return
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = INDENT_2CH
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.name, r1.font.size = FONT_SONG, SIZE_XIAOSI
            r1.font.bold = True
            r2 = p.add_run(text)
            r2.font.name, r2.font.size = FONT_SONG, SIZE_XIAOSI
        else:
            r = p.add_run(text)
            r.font.name, r.font.size = FONT_SONG, SIZE_XIAOSI

    # 一级「第一章」：黑体三号居中
    def add_heading1(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run(text)
        r.font.name, r.font.size = FONT_HEI, SIZE_SANHAO
        r.font.bold = True

    # 二级「一、」：黑体四号左对齐
    def add_heading2(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        r = p.add_run(text)
        r.font.name, r.font.size = FONT_HEI, SIZE_SIHAO
        r.font.bold = True

    # 三级「（一）」：楷体小四左对齐
    def add_heading3(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        r = p.add_run(text)
        r.font.name, r.font.size = FONT_KAI, SIZE_XIAOSI
        r.font.bold = False

    def add_chapter_line(doc, line):
        line = (line or "").strip()
        if not line:
            add_body(doc, "")
            return
        if re.match(r"^第[一二三四五六七八九十\d]+章\s*", line) or re.match(r"^第\s*\d+\s*章\s*", line):
            add_heading1(doc, line)
        elif re.match(r"^[一二三四五六七八九十]+[、．.]\s*", line) or re.match(r"^\d+[、．.]\s*", line):
            add_heading2(doc, line)
        elif re.match(r"^[（(][一二三四五六七八九十]+[)）]\s*", line) or re.match(r"^[（(]\d+[)）]\s*", line):
            add_heading3(doc, line)
        else:
            add_body(doc, line)

    add_heading1(doc, plain(book.title or "未命名"))
    doc.add_paragraph()
    if book.core_concept and (book.core_concept or "").strip():
        add_body(doc, plain(book.core_concept or ""), bold_prefix="核心构思：")
        doc.add_paragraph()
    preface = getattr(book, "preface", None) or ""
    if preface and preface.strip():
        add_heading1(doc, "前言")
        for line in plain(preface).splitlines():
            add_body(doc, (line or "").strip())
        doc.add_paragraph()
    if outline and outline.content and (outline.content or "").strip():
        add_heading1(doc, "全书大纲")
        for line in plain(outline.content or "").splitlines():
            add_body(doc, (line or "").strip())
        doc.add_paragraph()
    for ch in chapters:
        add_heading1(doc, "第 {} 章 {}".format(ch.order_index, plain(ch.title or "")))
        if ch.outline_content and (ch.outline_content or "").strip():
            add_body(doc, plain(ch.outline_content or ""), bold_prefix="本章要点：")
        ap = getattr(ch, "approved_sections", None)
        if isinstance(ap, str):
            try:
                ap = json.loads(ap) if ap.strip() else None
            except Exception:
                ap = None
        items = ap.get("items") if isinstance(ap, dict) else None
        body_src = ch.content or ""
        draft_src = getattr(ch, "draft_content", None) or ""
        if items and isinstance(items, dict) and draft_src.strip():
            body_src = draft_src
            for k, v in items.items():
                if k and v and str(v).strip():
                    body_src = _replace_markdown_section_by_title(body_src, str(k), str(v))
        elif not (body_src or "").strip() and draft_src.strip():
            body_src = draft_src
        if body_src and (body_src or "").strip():
            for line in plain(body_src or "").splitlines():
                add_chapter_line(doc, line)
        else:
            add_body(doc, "（本章正文尚未生成）")
        doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@app.route("/api/books/<int:bid>/generate-outline", methods=["POST"])
def generate_outline(bid):
    data = request.get_json() or {}
    revision_instruction = (data.get("revision_instruction") or "").strip() or None
    partial_revision = data.get("partial_revision") is True
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        if partial_revision:
            if not revision_instruction:
                return jsonify({"error": "仅局部修改时请填写修改意图"}), 400
            if not book.outline or not (book.outline.content or "").strip():
                return jsonify({"error": "仅局部修改需要先有全书大纲，请先生成大纲后再试"}), 400
        # 优先用请求体中的 content_type（与当前下拉框一致），避免与 DB 不同步导致仍走学术逻辑
        ct = (data.get("content_type") or getattr(book, "content_type", None) or "").strip().lower() or "academic"
        if ct not in ("academic", "novel"):
            ct = "academic"
        if hasattr(book, "content_type"):
            book.content_type = ct
        if hasattr(book, "book_type"):
            book.book_type = ct
        params = {"content_type": ct}
        if revision_instruction:
            params["revision_instruction"] = revision_instruction
        if partial_revision:
            params["partial_revision"] = True
        print("[generate-outline] book_id=%s content_type=%s (from_request=%s)" % (bid, ct, data.get("content_type")))
        task = GenerationTask(book_id=bid, task_type=TaskType.OUTLINE_LEVEL1, status=TaskStatus.PENDING, params=params)
        session.add(task)
        session.commit()
        session.refresh(task)
        tid = task.id
    except Exception as e:
        session.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()
    threading.Thread(target=_run_task_in_background, args=(bid, tid), daemon=True).start()
    return jsonify({"ok": True, "task_id": tid}), 202


@app.route("/api/books/<int:bid>/generate-preface", methods=["GET", "POST", "OPTIONS"], strict_slashes=False)
@app.route("/api/books/<int:bid>/generate-preface/", methods=["GET", "POST", "OPTIONS"], strict_slashes=False)
def generate_preface(bid):
    """生成前言（约 3000 字），后台任务。仅接受 POST；GET 返回 JSON 说明避免浏览器显示 HTML 405。"""
    if request.method == "OPTIONS":
        return "", 200, {"Allow": "POST", "Access-Control-Allow-Methods": "POST, OPTIONS"}
    if request.method != "POST":
        return jsonify({"error": "此接口仅支持 POST，请通过页面「生成前言」按钮操作"}), 405, {"Allow": "POST", "Content-Type": "application/json"}
    data = request.get_json() or {}
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        ct = (data.get("content_type") or getattr(book, "content_type", None) or "").strip().lower() or "academic"
        if ct not in ("academic", "novel"):
            ct = "academic"
        if hasattr(book, "content_type"):
            book.content_type = ct
        if hasattr(book, "book_type"):
            book.book_type = ct
        params = {"content_type": ct}
        task = GenerationTask(book_id=bid, task_type=TaskType.PREFACE, status=TaskStatus.PENDING, params=params)
        session.add(task)
        session.commit()
        session.refresh(task)
        tid = task.id
    except Exception as e:
        session.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()
    threading.Thread(target=_run_task_in_background, args=(bid, tid), daemon=True).start()
    return jsonify({"ok": True, "task_id": tid}), 202


@app.route("/api/books/<int:bid>/tasks/<int:tid>/cancel", methods=["POST"])
def cancel_task(bid, tid):
    """将运行中或等待中的任务标记为取消；引擎在下次检查时会中止并不写入结果。"""
    session = get_session()
    try:
        task = session.get(GenerationTask, tid)
        if not task or task.book_id != bid:
            return jsonify({"error": "任务不存在"}), 404
        if task.status not in (TaskStatus.PENDING, TaskStatus.RUNNING):
            return jsonify({"error": "任务已结束，无法取消"}), 400
        task.status = TaskStatus.CANCELLED
        session.commit()
        return jsonify({"ok": True})
    finally:
        session.close()


@app.route("/api/books/<int:bid>/chapters/from-outline", methods=["POST"])
def chapters_from_outline(bid):
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        outline = book.outline
        if not outline or not outline.raw_json:
            return jsonify({"error": "请先生成大纲"}), 400
        try:
            arr = json.loads(outline.raw_json)
        except Exception:
            return jsonify({"error": "大纲 JSON 解析失败"}), 400
        intro, fragments = _parse_outline_to_intro_and_fragments(outline.content or "")
        if getattr(outline, "intro", None) is not None:
            outline.intro = intro
        frag_by_idx = {idx: frag for idx, frag in fragments}
        created = []
        for i, item in enumerate(arr):
            idx = item.get("chapter_index", i + 1)
            title = item.get("title") or "第{}章".format(idx)
            frag = frag_by_idx.get(idx) or ""
            ch = Chapter(book_id=bid, order_index=idx, title=title)
            if hasattr(ch, "outline_fragment"):
                ch.outline_fragment = frag
            session.add(ch)
            session.flush()
            created.append({"id": ch.id, "order_index": ch.order_index, "title": ch.title})
        session.commit()
        return jsonify({"ok": True, "chapters": created})
    finally:
        session.close()


def _load_glossary(session, book_id):
    rows = session.execute(select(Term).where(Term.book_id == book_id).order_by(Term.id.asc())).scalars().all()
    return [{"term": r.term, "definition": r.definition or ""} for r in rows]


@app.route("/api/books/<int:bid>/chapters/<int:cid>/generate", methods=["POST"])
def generate_chapter(bid, cid):
    data = request.get_json() or {}
    return _enqueue_chapter_task(bid, cid, data)


def _enqueue_chapter_task(bid: int, cid: int, data: dict) -> tuple:
    """创建 CHAPTER 任务（可通过 data.stage 控制 chapter_draft / section_finalize）。"""
    revision_instruction = (data.get("revision_instruction") or "").strip() or None
    stage = (data.get("stage") or "").strip().lower() or None
    session = get_session()
    try:
        book = session.get(Book, bid)
        ch = session.get(Chapter, cid)
        if not book or not ch or ch.book_id != bid:
            return jsonify({"error": "书籍或章节不存在"}), 404
        # 优先用请求体中的 content_type（与当前下拉框一致）
        ct = (data.get("content_type") or getattr(book, "content_type", None) or "").strip().lower() or "academic"
        if ct not in ("academic", "novel"):
            ct = "academic"
        if hasattr(book, "content_type"):
            book.content_type = ct
        if hasattr(book, "book_type"):
            book.book_type = ct
        params = {"content_type": ct}
        if revision_instruction:
            params["revision_instruction"] = revision_instruction
        if stage:
            params["stage"] = stage
        print("[generate-chapter] book_id=%s chapter_id=%s content_type=%s (from_request=%s)" % (bid, cid, ct, data.get("content_type")))
        task = GenerationTask(book_id=bid, chapter_id=cid, task_type=TaskType.CHAPTER, status=TaskStatus.PENDING, params=params)
        session.add(task)
        session.commit()
        session.refresh(task)
        tid = task.id
    except Exception as e:
        session.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()
    threading.Thread(target=_run_task_in_background, args=(bid, tid), daemon=True).start()
    return jsonify({"ok": True, "task_id": tid}), 202


@app.route("/api/books/<int:bid>/chapters/<int:cid>/generate-draft", methods=["POST"])
def generate_chapter_draft(bid, cid):
    """章节草稿（约 3000 字）：stage=chapter_draft，后台任务。"""
    data = request.get_json() or {}
    data["stage"] = "chapter_draft"
    return _enqueue_chapter_task(bid, cid, data)


@app.route("/api/books/<int:bid>/chapters/<int:cid>/finalize-sections", methods=["POST"])
def finalize_chapter_sections(bid, cid):
    """逐节终稿：stage=section_finalize，每节约 3000 字，后台任务。"""
    data = request.get_json() or {}
    data["stage"] = "section_finalize"
    # 可选：只生成指定小节
    if data.get("section_title") is not None:
        data["section_title"] = (data.get("section_title") or "").strip()
    return _enqueue_chapter_task(bid, cid, data)


@app.route("/api/books/<int:bid>/chapters/<int:cid>/draft", methods=["GET"])
def get_chapter_draft(bid: int, cid: int):
    """获取该章保存的章节草稿（chapters.draft_content），用于查看对照。"""
    session = get_session()
    try:
        ch = session.get(Chapter, cid)
        if not ch or ch.book_id != bid:
            return jsonify({"error": "章节不存在"}), 404
        content = getattr(ch, "draft_content", None) or ""
        if not content.strip():
            return jsonify({"ok": True, "has_draft": False, "content": ""})
        return jsonify({"ok": True, "has_draft": True, "content": content[:200000]})
    finally:
        session.close()


@app.route("/api/books/<int:bid>/chapters/<int:cid>/rewrite", methods=["POST"])
def rewrite_chapter_slice(bid, cid):
    """
    定向重写：仅修改指定片段（如某一小节），不改动其他内容。
    Body:
      - scope: "chapter_section" | "chapter_paragraph_range" | "outline_chapter_fragment"
      - instruction: 修改意图（必填）
      - section_title: (scope=chapter_section) 小节标题（不含 ##）
      - anchor_start / anchor_end: (scope=chapter_paragraph_range) 用于定位段落的起止片段
    """
    data = request.get_json() or {}
    scope = (data.get("scope") or "").strip()
    instruction = (data.get("instruction") or "").strip()
    target = (data.get("target") or "").strip().lower() or "final"
    if target not in ("final", "draft"):
        target = "final"
    if not scope:
        return jsonify({"error": "缺少 scope"}), 400
    if not instruction:
        return jsonify({"error": "缺少 instruction（修改意图）"}), 400
    if scope == "chapter_section":
        if not (data.get("section_title") or "").strip():
            return jsonify({"error": "scope=chapter_section 时需提供 section_title"}), 400
    if scope == "chapter_paragraph_range":
        if not (data.get("anchor_start") or "").strip():
            return jsonify({"error": "scope=chapter_paragraph_range 时需提供 anchor_start"}), 400
        if not (data.get("anchor_end") or "").strip():
            return jsonify({"error": "scope=chapter_paragraph_range 时需提供 anchor_end"}), 400

    session = get_session()
    try:
        book = session.get(Book, bid)
        ch = session.get(Chapter, cid)
        if not book or not ch or ch.book_id != bid:
            return jsonify({"error": "书籍或章节不存在"}), 404
        ct = (data.get("content_type") or getattr(book, "content_type", None) or "").strip().lower() or "academic"
        if ct not in ("academic", "novel"):
            ct = "academic"
        if hasattr(book, "content_type"):
            book.content_type = ct
        if hasattr(book, "book_type"):
            book.book_type = ct
        params = {
            "content_type": ct,
            "scope": scope,
            "instruction": instruction,
            # 兼容旧逻辑：让引擎里仍能读取 revision_instruction
            "revision_instruction": instruction,
            "target": target,
        }
        if data.get("section_title"):
            params["section_title"] = (data.get("section_title") or "").strip()
        if data.get("anchor_start"):
            params["anchor_start"] = (data.get("anchor_start") or "").strip()
        if data.get("anchor_end"):
            params["anchor_end"] = (data.get("anchor_end") or "").strip()

        task = GenerationTask(book_id=bid, chapter_id=cid, task_type=TaskType.REWRITE, status=TaskStatus.PENDING, params=params)
        session.add(task)
        session.commit()
        session.refresh(task)
        tid = task.id
    except Exception as e:
        session.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()
    threading.Thread(target=_run_task_in_background, args=(bid, tid), daemon=True).start()
    return jsonify({"ok": True, "task_id": tid}), 202


@app.route("/api/books/<int:bid>/chapters/<int:cid>/draft", methods=["PUT"])
def update_chapter_draft(bid: int, cid: int):
    """保存章节草稿（draft_content），不影响终稿 content。"""
    data = request.get_json() or {}
    raw = data.get("draft_content")
    if raw is None:
        return jsonify({"error": "请提供 draft_content 字段"}), 400
    session = get_session()
    try:
        ch = session.get(Chapter, cid)
        if not ch or ch.book_id != bid:
            return jsonify({"error": "章节不存在"}), 404
        ch.draft_content = _normalize_paragraph_spacing(raw)
        session.commit()
        return jsonify({"ok": True})
    finally:
        session.close()


@app.route("/api/books/<int:bid>/chapters/<int:cid>/approve-section", methods=["POST"])
def approve_chapter_section(bid: int, cid: int):
    """标记某个小节审核通过（用于导出）。"""
    data = request.get_json() or {}
    title = (data.get("section_title") or "").strip()
    if not title:
        return jsonify({"error": "section_title 不能为空"}), 400
    session = get_session()
    try:
        ch = session.get(Chapter, cid)
        if not ch or ch.book_id != bid:
            return jsonify({"error": "章节不存在"}), 404
        cur = getattr(ch, "approved_sections", None)
        if isinstance(cur, str):
            try:
                cur = json.loads(cur) if cur.strip() else None
            except Exception:
                cur = None
        if not isinstance(cur, dict):
            cur = {}
        titles = cur.get("titles")
        if not isinstance(titles, list):
            titles = []
        if title not in titles:
            titles.append(title)
        cur["titles"] = titles
        items = cur.get("items")
        if not isinstance(items, dict):
            items = {}
        # 存储该小节的当前终稿内容（包含标题行），用于导出时替换草稿对应位置
        full = (ch.content or "").strip()
        section_text = ""
        if full:
            try:
                from core.engine import _split_markdown_section_by_title
                _, sec, _ = _split_markdown_section_by_title(full, title)
                section_text = sec
            except Exception:
                section_text = ""
        if section_text:
            items[title] = section_text
        cur["items"] = items
        ch.approved_sections = cur
        session.commit()
        return jsonify({"ok": True, "approved_sections": cur})
    finally:
        session.close()


# ---------- References / Citations ----------

@app.route("/api/books/<int:bid>/references", methods=["GET"])
def get_references(bid: int):
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        refs = list_references(session, book_id=bid)
        return jsonify({
            "ok": True,
            "items": [{
                "id": r.id,
                "citation_key": r.citation_key,
                "title": r.title or "",
                "file_path": r.file_path or "",
                "content_extract": (r.content_extract or "")[:8000],
                "meta": r.meta or {},
                "created_at": r.created_at.isoformat() if getattr(r, "created_at", None) else None,
            } for r in refs]
        })
    finally:
        session.close()


@app.route("/api/books/<int:bid>/references/from-url", methods=["POST"])
def add_reference_url(bid: int):
    data = request.get_json() or {}
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "url 不能为空"}), 400
    citation_key = (data.get("citation_key") or "").strip() or None
    title = (data.get("title") or "").strip() or None
    with_llm = data.get("with_llm_summary", True) is True
    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        ref = add_reference_from_url(
            session,
            book_id=bid,
            url=url,
            citation_key=citation_key,
            title=title,
            with_llm_summary=with_llm,
        )
        session.commit()
        return jsonify({"ok": True, "reference_id": ref.id})
    except Exception as e:
        session.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


@app.route("/api/books/<int:bid>/references/upload", methods=["POST"])
def upload_reference_file(bid: int):
    """
    上传本地文献文件（txt/md/html/pdf）。
    - 表单字段: file (required), citation_key (optional), title(optional), with_llm_summary(optional: true/false)
    """
    if "file" not in request.files:
        return jsonify({"error": "缺少 file"}), 400
    f = request.files["file"]
    if not f or not getattr(f, "filename", ""):
        return jsonify({"error": "文件为空"}), 400
    filename = _sanitize_filename(str(f.filename))
    save_path = REF_DIR / (datetime.now().strftime("%Y%m%d-%H%M%S-") + filename)
    f.save(str(save_path))
    citation_key = (request.form.get("citation_key") or "").strip() or None
    title = (request.form.get("title") or "").strip() or None
    with_llm = (request.form.get("with_llm_summary") or "").strip().lower()
    with_llm = False if with_llm in ("0", "false", "no") else True

    session = get_session()
    try:
        book = session.get(Book, bid)
        if not book:
            return jsonify({"error": "书籍不存在"}), 404
        # pdf 暂不解析正文（避免额外依赖）；仍保存 file_path，后续可加 pypdf
        ref = add_reference_from_file(
            session,
            book_id=bid,
            file_path=str(save_path),
            citation_key=citation_key,
            title=title,
            with_llm_summary=with_llm,
        )
        session.commit()
        return jsonify({"ok": True, "reference_id": ref.id, "file_path": str(save_path)})
    except Exception as e:
        session.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


@app.route("/api/books/<int:bid>/references/<int:rid>", methods=["DELETE"])
def delete_reference(bid: int, rid: int):
    session = get_session()
    try:
        ref = session.get(Reference, rid)
        if not ref or ref.book_id != bid:
            return jsonify({"error": "参考文献不存在"}), 404
        # 先删 citations，避免外键约束
        session.execute(delete(Citation).where(Citation.reference_id == rid))
        session.delete(ref)
        session.commit()
        return jsonify({"ok": True})
    except Exception as e:
        session.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


@app.route("/api/books/<int:bid>/tasks/current", methods=["GET"])
def current_task(bid):
    session = get_session()
    try:
        running = list(session.execute(
            select(GenerationTask).where(GenerationTask.book_id == bid, GenerationTask.status == TaskStatus.RUNNING).order_by(GenerationTask.id.desc()).limit(1)
        ).scalars().all())
        # 仅返回“最近失败”（避免历史失败长期占用 UI）
        cutoff = datetime.now(timezone.utc).replace(tzinfo=None) - timedelta(seconds=45)
        failed = list(session.execute(
            select(GenerationTask).where(
                GenerationTask.book_id == bid,
                GenerationTask.status == TaskStatus.FAILED,
                GenerationTask.completed_at != None,  # noqa: E711
                GenerationTask.completed_at >= cutoff,
            ).order_by(GenerationTask.id.desc()).limit(1)
        ).scalars().all())
        out = {}
        if running:
            t = running[0]
            out["running"] = {
                "id": t.id,
                "chapter_id": t.chapter_id,
                "task_type": t.task_type.value if hasattr(t.task_type, "value") else str(t.task_type),
                "progress_message": t.progress_message or "",
                "current_output": (t.current_output or "")[:50000],
            }
        # 仅在没有运行中任务时返回上次失败错误，避免旧错误盖住当前进度
        if failed and not running:
            out["last_error"] = failed[0].error_message or ""
        return jsonify(out)
    finally:
        session.close()


@app.route("/api/books/<int:bid>/tasks/stream", methods=["GET"])
def stream_task(bid: int):
    """
    SSE: 推送任务进度与 current_output 增量，前端可逐步展示（EventSource）。
    Query:
      - task_id (required): 要订阅的任务 ID

    Event data (JSON):
      { task_id, chapter_id, task_type, status, progress_message, delta, done }
    """
    task_id_raw = (request.args.get("task_id") or "").strip()
    try:
        tid = int(task_id_raw)
    except Exception:
        return jsonify({"error": "task_id 必须为整数"}), 400

    def _event(payload: dict) -> str:
        return "data: " + json.dumps(payload, ensure_ascii=False) + "\n\n"

    @stream_with_context
    def gen():
        last_len = 0
        idle_ticks = 0
        # 首包：让前端尽快进入“已连接”状态
        yield _event({"task_id": tid, "delta": "", "done": False})
        while True:
            session = get_session()
            try:
                task = session.get(GenerationTask, tid)
                if not task or task.book_id != bid:
                    yield _event({"task_id": tid, "done": True, "status": "missing", "error": "任务不存在"})
                    return
                status = task.status.value if hasattr(task.status, "value") else str(task.status)
                task_type = task.task_type.value if hasattr(task.task_type, "value") else str(task.task_type)
                progress = task.progress_message or ""
                cur = task.current_output or ""
                if len(cur) > last_len:
                    delta = cur[last_len:]
                    last_len = len(cur)
                    idle_ticks = 0
                    yield _event({
                        "task_id": tid,
                        "chapter_id": task.chapter_id,
                        "task_type": task_type,
                        "status": status,
                        "progress_message": progress,
                        "delta": delta,
                        "done": status in (TaskStatus.COMPLETED.value, TaskStatus.FAILED.value, TaskStatus.CANCELLED.value),
                    })
                else:
                    idle_ticks += 1
                    # 心跳：避免代理/浏览器认为连接空闲
                    if idle_ticks % 10 == 0:
                        yield _event({
                            "task_id": tid,
                            "chapter_id": task.chapter_id,
                            "task_type": task_type,
                            "status": status,
                            "progress_message": progress,
                            "delta": "",
                            "done": status in (TaskStatus.COMPLETED.value, TaskStatus.FAILED.value, TaskStatus.CANCELLED.value),
                        })

                if status in (TaskStatus.COMPLETED.value, TaskStatus.FAILED.value, TaskStatus.CANCELLED.value):
                    # 末包：确保 done=true
                    yield _event({
                        "task_id": tid,
                        "chapter_id": task.chapter_id,
                        "task_type": task_type,
                        "status": status,
                        "progress_message": progress,
                        "delta": "",
                        "done": True,
                    })
                    return
            finally:
                session.close()

            # 频率：0.35s 左右，足够“流式”且不压垮 SQLite
            import time
            time.sleep(0.35)

    headers = {
        "Content-Type": "text/event-stream; charset=utf-8",
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
        "X-Accel-Buffering": "no",  # nginx: disable buffering if present
    }
    return Response(gen(), headers=headers)


# ---------- 启动 ----------

if __name__ == "__main__":
    init_db()
    ensure_content_type_column()
    _cleanup_stale_running_tasks()
    print("自动化写书 已启动： http://127.0.0.1:5000")
    print("关闭窗口或 Ctrl+C 停止服务。")
    app.run(host="127.0.0.1", port=5000, debug=False, threaded=True)
